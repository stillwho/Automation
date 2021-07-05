import base64
import tkinter
from datetime import datetime
from tkinter import messagebox, filedialog
from docx.opc.oxml import qn
from xlutils.copy import copy
import time
from tkinter import *
from random import choice
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from openpyxl import load_workbook
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt  # 导入单位换算函数
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

'''
文本转移
'''

import sys
import os

'''
界址点表
'''
import xlrd
import openpyxl
from openpyxl.styles import Border, Side, colors
from openpyxl.styles import Alignment
from openpyxl.styles import Font
from openpyxl import Workbook


# 用于将输出进行重定向
class TextRedirector(object):
    def __init__(self, widget, tag="stdout"):
        self.widget = widget
        self.tag = tag

    def write(self, str):
        self.widget.configure(state="normal")
        self.widget.insert("end", str, (self.tag,))
        self.widget.see('end')
        self.widget.configure(state="disabled")

# 主程序
class Aotumation(tkinter.Tk):
    def __init__(self):

        # 创建根窗口
        self.root = Tk()
        # 根窗口主标题
        self.root.title('Aoanauto_v1.0.0')
        # 根窗口大小
        self.root.geometry('550x650')
        # 窗口图标，用ASCII表示图标，防止pyinstall无法打包
        self.icon = open("icon.ico", "wb+")
        self.icon.write(base64.b64decode(
            'AAABAAEAICAAAAEAIACoEAAAFgAAACgAAAAgAAAAQAAAAAEAIAAAAAAAABAAAAAAAAAAAAAAAAAAAAAAAAD//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////f7+//X9/f/6/v7//v7+//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////7///7+/v/w/Pz/oe/1/7bw9v/6/f7//v7////+/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////v/+/v7/7vv9/6Hs8/+P5/L/jeXw/67q8v/4/f3//v7///7//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////v7+/+z7/P+e6fH/jeTu/4vi7P+J4Ov/h9zr/6fk7f/2/P3//v7///////////////////////////////////////////////////////////////////////////////////////////////////////////////////7+/v/u+/v/m+Xw/4vg7/+J3uz/mOHs/5bf6f+F2Ob/g9Xm/6He6f/2/P3//v7+///////////////////////////////////////////////////////////////////////+/v///v7///7///////////////7+/v/+/v7/7Pr6/5nj7/+J3ev/htvp/5re6//w+vz/7fn6/5LY5f9/0OH/f83g/5rX5P/1+/v//v7+//7+//////////////7+///+/v////////////////////////////////////////3+/v/t/fz//v7+//7+///+//7//f7+/+z5+/+W4Oz/iNrp/4XY5/+a3ej/8fr7//7+/v/9/v7/7vj6/4/S4f97yNv/esba/5bP4P/y+fr//v7+/////v/+/v7//f7+/+z8/P/9/v7//////////////////////////////////v7+/835+//K+Pv//v7+//3+/v/q+Pr/k97r/4XX5/+D1eX/mtvo//H6+//+/v7////////////+/v7/7vj5/47M3P94wNf/dr/U/47I2//v9/r//v7+//3+/v/F9fn/z/f6//7+/v/////////////////////////////////+/v//9f39/6Tx9f/3/f3/6fj5/5Da6P+C1OX/gdLk/5vZ5v/y+vv////+/////////////////////v/+//7/7/j5/4zG2f9xudH/cbfP/4jB1f/u9vj/9fz9/6Lt9P/3/f7//v7///////////////////////////////////7////+/v7/x/T4/7zv9f+Q1+b/gdHi/4DO4f+Z1+T/9Pv7///+/v/////////////////////////////////+/v7/8Pf5/4jB1P9sscz/bK/K/4C4z/+56fH/y/P3//7+/v/+/////////////////////////////////////v7+//7+/v/e9Pf/jePu/37P4P98zN3/m9bk//X7/P/+/v7////////////////////////////////////////////+/v7/8fj5/4i80f9oqsX/aavI/4jb6f/i8fT//v7+//7//v////////////////////////////7//v/9/v7/5PX3/4jQ4f+G2Ob/hNPj/5vV4//2+/z//v7+//7+/v/6+vr/6enp/9zd3f/X2Nj/2NjY/93d3f/p6en/+vr6//7+/v/+/v7/8vj4/4a3zv91wdX/d8LW/3Sqxv/n8PT//v7+//7+/v///////////////////////v/+/+T09v+Gzt7/e8jb/3zI2/+O3Oj/8/v8//7////5+fn/2dnZ/9PT0//j4+P/q62t/6yurv+hoqL/qaiq/+Pj4//S0tL/29vb//r6+v/+//7/7vf5/4fU4/9lpMP/Yp6//3Soxf/j7/P//v7+///////////////////////u9/j/ltHf/4vL3P+Jytv/ptfj/7Xn8P/R8PT/8/Pz/9DQ0P/w8PD//v7+/+np6f+xs7P/Q0VK/0hKTv+Gh4j/z8/R//7+/v/u7u7/z8/P//b29v/K7fP/uubv/5O90f91qsT/darE/4Ozyf/t8/b///////////////////////7+/v/+/v7//v7+//3+/v/9/v7/8Pr7/5zc5//d39//3d3d//7+/v//////4+Pj/76+v/8ZGSD/Ghsg/4qLjf+/v8D///////7+/v/a2tr/3+Li/5ra5v/0+/z//v7+//3+/v/9/v7//f79//3+/v/+/v7////////////////////////////////////////////9/v7/tuLt/8zq7//V1tb/0tLS/+7u7v/19fX/1tfW/7Gxs/+anJ7/ODo9/+zs7P/t7e3/0tLS/9PU0//H5+3/u+Xt//7+/v////////////////////////////////////////////////////////////////////////////7//v/t+Pr/ltXj//Dy8//e3t7/q6ur/8vLy//MzMz/y8vL/6amp/+trq//wsLC/7Gxsf/Ozs7/5Ofn/5XT4f/x+fv//v7+//////////////////////////////////////////////////////////////////////////////////7+/v+u3Of/vd3l/8bGx//T09T/iIqL/2xucP+ur7D/YmNo/4SFiP+Mjo//ra+w/8vNzf+qzNX/s97p//7+/v///////////////////////////////////////////////////////////////////////////////////////v/+/+n2+P+Hydv/6/b4/9jZ2v/a2tv/qKmp/9fY2f99fYD/6erq/5iamv/d3t//4e7x/4nI3P/t9/r//v7+//////////////////////////////////////////////////////////////////////////////////////////7//v7+/6XV4v+Bw9j/x+Tr/+Tm5v/7+/z/+/v7/7u7vP/+/v7/4uTl/8Tj6/+Bw9n/rNnk//3+/v///v//////////////////////////////////////////////////////////////////////////////////////////////////5fP1/3S70f9yudH/2ezx//7+/v///////Pz8//7+/v/U6/D/c73U/3i91P/s9fj////////////////////////////////////////////////////////////////////////////////////////////////////////////9/v7/nc7d/2+1zf+Mw9b//P3+//7+/v/+/v//+/39/4nD1v9wuND/p9Lg//7+/v////////////////////////////////////////////////////////////////////////////////////////////////////////////7+/v/i8PT/brHM/26wy//F4Oj//v7+//7+/v/B3en/b7TP/2+1zv/o8/b//v7+//////////////////////////////////////////////////////////////////////////////////////////////////////////////////3+/v+Yxdj/aq3I/3m0zf/0+vv/8vj6/3i3z/9qsMv/nsvc//3+/v///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////v7+/97s8f9pqcb/a7HL/6vk7f+p4+3/brPN/22syf/k8PT//v7+/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////v7+/5PD1v+G2uf/jufw/43n7/+L3+z/msfZ//7+/v/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+/v//w+bu/4LR4f+G2Oj/gNLi/4TU5P/Q8PP//v7///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////3+/v/L5u3/dL7U/3K80v9tt83/eMTZ/9zw8//+/v7//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////v7+//j8/P+PwdX/a67K/2ywyv+ezt3//P79//7+/v////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////v9/P/a6e7/3Ovw//z+/f//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA='))  # tmp.write(base64.b64decode('粘贴icon2.py字符串内容'))
        self.icon.close()
        self.root.iconbitmap('icon.ico')
        # 框架分布
        self.frame1 = Frame(self.root)
        self.frame2 = Frame(self.root)
        self.frame3 = Frame(self.root)
        # 创建菜单-----start--------------------------------------------------
        self.menubar = Menu(self.root)
        helpmenu = Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label='帮助', menu=helpmenu)
        helpmenu.add_command(label='快速引导', command=self.help_quickstart)
        helpmenu.add_command(label='用户手册', command=self.help_help)
        helpmenu.add_separator()
        helpmenu.add_command(label='技术支持', command=self.help_support)
        helpmenu.add_command(label='定制服务', command=self.help_service)
        # 可以添加其他工具
        # toolmenu = Menu(self.menubar, tearoff=0)
        # self.menubar.add_cascade(label='工具')
        # toolmenu.add_command(label='正在开发')
        aboutmenu = Menu(self.menubar, tearoff=0)
        self.menubar.add_cascade(label='关于', menu=aboutmenu)
        aboutmenu.add_command(label='作者', command=self.about_author)
        aboutmenu.add_command(label='软件', command=self.about_software)
        aboutmenu.add_command(label='加入我们', command=self.about_joinus)

        self.root.config(menu=self.menubar)
        # 创建菜单-----end--------------------------------------------------

        # 生成主界面其余widget-----start--------------------------------------
        # 文本框，按列表排列
        self.projectLeader = Label(self.frame2, text='项目负责人：').grid(row=1, column=0)
        self.reportYear = Label(self.frame2, text='年份：').grid(row=2, column=0)
        self.number = Label(self.frame2, text='号：').grid(row=3, column=0)
        self.entrustingParty = Label(self.frame2, text='委托方：').grid(row=4, column=0)
        self.reportTime = Label(self.frame2, text='报告时间：').grid(row=1, column=2)
        self.measuringTime = Label(self.frame2, text='测量时间：').grid(row=2, column=2)
        self.scale = Label(self.frame2, text='比例：').grid(row=3, column=2)
        self.database = Label(self.frame2, text='数据库图：').grid(row=8, column=0)
        self.striograph = Label(self.frame2, text='影像图：').grid(row=9, column=0)
        self.programme = Label(self.frame2, text='规划图：').grid(row=10, column=0)

        self.inputLandFormFilePath = Label(self.frame2, text='地类文件：').grid(row=11, column=0)
        self.outFilePath = Label(self.frame2, text='保存路径：').grid(row=13, column=0)
        self.inputJZDFormFilePath = Label(self.frame2, text='界址点文件：').grid(row=12, column=0)
        self.standbyPicture = Label(self.frame2, text='备用图：').grid(row=14, column=0)

        # 复选框
        self.var1 = IntVar()
        self.check17 = Radiobutton(self.frame2, text='资质延期图', variable=self.var1, value=1,
                                   command=self.print_selection).grid(row=4, column=2)
        # Entry默认值
        default_year = StringVar(self.root, value=datetime.now().year)
        default_date = StringVar(self.root,
                                 value=time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='日'))
        default_scale = StringVar(self.root, value='1:500')

        # 输入框排布
        self.e_projectLeader = Entry(self.frame2)
        self.e_reportYear = Entry(self.frame2, textvariable=default_year)
        self.e_number = Entry(self.frame2)
        self.e_entrustingParty = Entry(self.frame2)
        self.e_reportTime = Entry(self.frame2, textvariable=default_date)
        self.e_measuringTime = Entry(self.frame2, textvariable=default_date)
        self.e_scale = Entry(self.frame2, textvariable=default_scale)
        self.e_database = Entry(self.frame2)
        self.e_striograph = Entry(self.frame2)
        self.e_programme = Entry(self.frame2)
        self.e_inputLandFormFilePath = Entry(self.frame2)
        self.e_inputJZDFormFilePath = Entry(self.frame2)
        self.e_outFilePath = Entry(self.frame2)
        self.e_standbyPicture = Entry(self.frame2)

        self.e_projectLeader.grid(row=1, column=1, padx=10, pady=5)
        self.e_reportYear.grid(row=2, column=1, padx=10, pady=5)
        self.e_number.grid(row=3, column=1, padx=10, pady=5)
        self.e_entrustingParty.grid(row=4, column=1, padx=10, pady=5)
        self.e_reportTime.grid(row=1, column=4, padx=10, pady=5)
        self.e_measuringTime.grid(row=2, column=4, padx=10, pady=5)
        self.e_scale.grid(row=3, column=4, padx=10, pady=5)
        self.e_database.grid(row=8, column=1, padx=20, pady=5)
        self.e_striograph.grid(row=9, column=1, padx=10, pady=5)
        self.e_programme.grid(row=10, column=1, padx=10, pady=5)
        self.e_inputLandFormFilePath.grid(row=11, column=1, padx=10, pady=5)
        self.e_inputJZDFormFilePath.grid(row=12, column=1, padx=10, pady=5)
        self.e_outFilePath.grid(row=13, column=1, padx=10, pady=5)
        self.e_standbyPicture.grid(row=14, column=1, padx=10, pady=5)

        # 创建按钮
        # 调用函数
        self.add_button = Button(self.frame3, text='执行', width=10,
                                 command=lambda: [self.add_context(), self.change_docx()]) \
            .grid(row=3, column=2, sticky=W, padx=10, pady=5)
        # 直接退出
        self.quit_button = Button(self.frame3, text='退出', width=10, command=self.root.quit) \
            .grid(row=3, column=3, sticky=E, padx=10, pady=5)
        # 选择输入文件
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.openfile(self.e_database)) \
            .grid(row=8, column=2, padx=10, pady=5)
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.openfile(self.e_striograph)) \
            .grid(row=9, column=2, padx=10, pady=5)
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.openfile(self.e_programme)) \
            .grid(row=10, column=2, padx=10, pady=5)
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.openfile(self.e_inputLandFormFilePath)) \
            .grid(row=11, column=2, padx=10, pady=5)
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.openfile(self.e_inputJZDFormFilePath)) \
            .grid(row=12, column=2, padx=10, pady=5)
        # 选择存储路径
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.savefile(self.e_outFilePath)) \
            .grid(row=13, column=2, padx=10, pady=5)
        # 选择选择输入文件
        self.outFilePath = Button(self.frame2, text='浏览', width=10, height='1',
                                  command=lambda: self.openfile(self.e_standbyPicture)) \
            .grid(row=14, column=2, padx=10, pady=5)

        # 通用输出
        self.lable15 = Label(self.frame3, text='运行状态： ').grid(row=4, column=0)
        self.console = Text(self.frame3, width='40', height='10', state='disabled')
        self.console.config(bg='#eee', fg='#888')
        self.console.tag_configure("stderr", foreground="#b22222")
        self.console.grid(row=4, column=1, columnspan=3, pady=3)
        sys.stdout = TextRedirector(self.console, "stdout")
        sys.stderr = TextRedirector(self.console, "stderr")
        print('欢迎使用Aoanauto...\n在使用前请先阅读帮助中的用户手册...\n然后填写上述内容...\n点击执行即可')

    # 分布包装
    def gui_arrang(self):
        self.frame1.pack()
        self.frame2.pack(padx=10, pady=10)
        self.frame3.pack()

    # 获取数据
    def add_context(self):

        # 打开文件
        self.tem_excel = xlrd.open_workbook('报告信息.xls')
        self.tem_sheet = self.tem_excel.sheet_by_index(0)

        self.new_excel = copy(self.tem_excel)
        self.new_sheet = self.new_excel.get_sheet(0)

        # 可修改写入数据
        self.new_sheet.write(1, 1, self.e_projectLeader.get())
        self.new_sheet.write(1, 2, self.e_reportYear.get())
        self.new_sheet.write(1, 3, self.e_number.get())
        self.new_sheet.write(1, 4, self.e_entrustingParty.get())
        self.new_sheet.write(1, 6, self.e_reportTime.get())
        self.new_sheet.write(1, 7, self.e_measuringTime.get())
        self.new_sheet.write(1, 8, self.e_scale.get())
        try:
            if self.e_inputLandFormFilePath.get() != '':
                self.new_sheet.write(1, 0, self.print_projectname())
                self.new_sheet.write(1, 5, self.print_landform_totalArea())
                self.new_sheet.write(1, 9, self.landform_region())
        except Exception as e:
            sys.stderr.write('请选择土地分类面积表...')
            sys.stderr.write(e)

        if os.path.exists('报告信息.xls'):
            os.remove('报告信息.xls')

        self.new_excel.save('报告信息.xls')

    # 数据替换
    def change_docx(self):
        self.document = Document('报告模板.docx')
        # 日期设定为当日
        # today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='日')
        # 打开并扫描表格，替换
        xlsx = xlrd.open_workbook('报告信息.xls')
        sheet = xlsx.sheet_by_index(0)
        for table_row in range(1, sheet.nrows):
            for table_col in range(0, sheet.ncols):
                self.change_text(str(sheet.cell_value(0, table_col)), str(sheet.cell_value(table_row, table_col)))
            # 可修改条件
            # self.change_text('xlandformRegionx', today)

        # 写入界址点成果表
        if self.e_inputJZDFormFilePath.get() != '':
            # 修改导入的.xls文件
            self.creat_JZDForm()
            # 写入word
            wb = load_workbook('test.xlsx', data_only=True)
            self.document.styles['Normal'].font.name = u'宋体'
            self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            self.document.styles['Normal'].font.size = Pt(12)

            font = self.document.add_paragraph().add_run('界址点成果表').font
            font.size = Pt(22)
            font.bold = True

            self.table_styles = ['Table Grid']

            for ws in wb.worksheets:
                self.section = self.document.sections[0]
                # 设置A4纸的宽度
                self.section.page_width = Cm(21)
                # 设置A4纸的高度
                self.section.page_height = Cm(29.7)
                # 创建表头
                self.rows = list(ws.rows)

                self.document.add_paragraph('')
                self.document.add_paragraph('    地块{0}由点号J1-J{1},J1围成'.format(ws.title[5], len(self.rows) - 1))
                self.document.add_paragraph(
                    '    占地面积：P={:.2f}平方米，合{}亩。'.format(float(ws.cell(2, 4).value), ws.cell(2, 5).value))
                # 建表，通过读取文件直接生成相应数量的空表
                self.table = self.document.add_table(rows=len(self.rows),
                                                     cols=len(self.rows[0]) - 2,
                                                     style=choice(self.table_styles))
                # 写入数据并调整保留小数点位数
                for table_col in range(1, 4):
                    self.table.cell(0, table_col - 1).text = str(ws.cell(1, table_col).value)
                for table_row in range(2, len(self.rows) + 1):
                    for table_col in range(2, 4):
                        self.table.cell(table_row - 1, table_col - 1).text = format(ws.cell(table_row, table_col).value,
                                                                                    '.3f')
                for table_row in range(2, len(self.rows) + 1):
                    self.table.cell(table_row - 1, 0).text = str(ws.cell(table_row, 1).value)
                for table_row in range(0, len(self.rows)):
                    for table_col in range(0, 3):
                        # 垂直居中
                        self.table.cell(table_row, table_col).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                        self.table.cell(table_row, table_col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for row in self.table.rows:
                    # 设置行高
                    row.height = Cm(1)
        # 插入分页符
        self.document.add_page_break()
        # 生成地类分类表文件
        if self.e_inputLandFormFilePath.get() != '':
            self.landform_creat()
            self.landform_region()
            self.print_landform_totalArea()

        # 如果Entey有值则添加图片
        # 添加数据库图片
        if self.e_database.get() != '':
            try:
                self.document.add_paragraph('数据库')
                self.document.add_picture(self.e_database.get(), width=Inches(6))
            except Exception as e:
                sys.stderr.write('没有找到数据库文件...')
                sys.stderr.write(e)
        # 添加影像图片
        if self.e_striograph.get() != '':
            try:
                self.document.add_paragraph('影像图')
                self.document.add_picture(self.e_striograph.get(), width=Inches(6))
            except Exception as e:
                sys.stderr.write('没有找到影像图文件...')
                sys.stderr.write(e)
        # 添加规划图片
        if self.e_programme.get() != '':
            try:
                self.document.add_paragraph('规划图')
                self.document.add_picture(self.e_programme.get(), width=Inches(6))
            except Exception as e:
                sys.stderr.write('没有找到规划图文件...')
                sys.stderr.write(e)
        # 添加备用图片
        if self.e_standbyPicture.get() != '':
            try:
                self.document.add_picture(self.e_standbyPicture.get(), width=Inches(6))
            except Exception as e:
                sys.stderr.write('没有找到备用图文件...')
                sys.stderr.write(e)
        # 添加资质图片
        try:
            self.document.add_picture('测绘资质.jpg', width=Inches(6))
        except Exception as e:
            sys.stderr.write('没有找到资质文件...')
            sys.stderr.write(e)
        # 添加延期图片
        if self.var1.get() == 1:
            try:
                self.document.add_picture('延期图.jpg', width=Inches(6))
            except Exception as e:
                sys.stderr.write('没有找到延期图文件...')
                sys.stderr.write(e)
        # 保存文件路径
        try:
            self.document.save('%s' % self.e_outFilePath.get())
            print('%s完成' % self.e_outFilePath.get())
        except Exception as e:
            sys.stderr.write('请填写保存文件路径...')
            sys.stderr.write(e)

    # 扫描文本和表格并替换
    def change_text(self, old_text, new_text):
        all_paragraphs = self.document.paragraphs
        for paragraph in all_paragraphs:
            for run in paragraph.runs:
                run_text = run.text.replace(old_text, new_text)
                run.text = run_text

        all_tables = self.document.tables
        for table in all_tables:
            for row in table.rows:
                for cell in row.cells:
                    all_cell_paragraphs = cell.paragraphs
                    for cell_paragraph in all_cell_paragraphs:
                        for cell_run in cell_paragraph.runs:
                            cell_run_text = cell_run.text.replace(old_text, new_text)
                            cell_run.text = cell_run_text

    def about_author(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('关于开发者...')
        self.root.showinfo(title='关于开发者', message='发起人：敖岸测绘有限公司\n开发团队：和山数据有限公司')

    def about_software(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('关于Aoanauto...')
        messagebox.showinfo(title='关于Aoanauto',
                            message='版本：v1.0.0\n核心功能：自动化制作测绘相关报告\n')

    def about_joinus(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('加入我们...')
        messagebox.showinfo(title='加入我们',
                            message='想要加入我们吗？发送简历和项目到邮箱吧！\n联系方式:18970937294\n')

    def help_quickstart(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('快速引导...')
        messagebox.showinfo(title='快速引导',
                            message='1.选择arcgis中导出的地类表和界址点表;\n2.在主界面填写报告信息;\n3.点击各界面中的帮助按钮会在运行状态栏输出工具使用指南')

    def help_help(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('打开帮助文件...')
        try:
            os.startfile('help.pdf')
        except Exception as e:
            sys.stderr.write('help.pdf文件缺失...')
            sys.stderr.write(e)

    def help_support(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('发送邮件或加入用户群获取技术支持...')
        messagebox.showinfo(title='技术支持', message='发送邮件至heshandata@163.com')

    def help_service(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('定制服务...')
        messagebox.showinfo(title='定制服务', message='联系方式:18970937294')

    #选择保存路径
    def savefile(self, widget):
        str_outFilePath = filedialog.asksaveasfilename()
        try:
            if str_outFilePath == '':
                str_outFilePath = '{}报告.docx'.format(self.print_projectname())
            widget.delete(0, 'end')
            widget.insert('end', str_outFilePath)
            self.console.config(state='normal')
            self.console.delete(1.0, 'end')
            self.console.config(state='disable')
            print('设置输出文件为：' + str_outFilePath)
        except Exception as e:
            sys.stderr.write('请选择土地分类面积表...')
            sys.stderr.write(e)

    # 选择输入文件
    def openfile(self, widget):
        str_inputFilePath = filedialog.askopenfilename()
        widget.delete(0, 'end')
        widget.insert('end', str_inputFilePath)
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('设置输入文件为：' + str_inputFilePath)

    def print_selection(self):
        self.console.config(state='normal')
        self.console.delete(1.0, 'end')
        self.console.config(state='disable')
        print('选取资质图')

    # 获取地区所有名称
    def landform_region(self):
        # 存储数据表QSDWMC中去重后的名称
        listRegion = []
        file = self.e_inputLandFormFilePath.get()
        # 创建工作空间
        workbook = xlrd.open_workbook(file)
        # 选取第0个表
        sheet1 = workbook.sheet_by_index(0)
        # 获取第5列数据
        iplist = (sheet1.col_values(5))
        # 把第一行标题弹出去
        iplist.pop(0)

        # 去掉重复的
        for i in iplist:
            if i not in listRegion:
                listRegion.append(i)
        if listRegion[-1] == '':
            listRegion.pop(-1)
        region = '、'.join(listRegion)

        return region

    #获取总面积
    def print_landform_totalArea(self):
        eachArea = 0
        file = self.e_inputLandFormFilePath.get()
        # 创建工作空间
        workbook = xlrd.open_workbook(file)
        # 选取第0个表
        sheet1 = workbook.sheet_by_index(0)
        eachArea_list = (sheet1.col_values(13))
        eachArea_list.pop(0)
        # eachArea_list.pop(-1)
        #去出空值
        eachArea_list = [i for i in eachArea_list if i != '']
        #算总面积
        for i in range(0, len(eachArea_list)):
            if eachArea_list[i] is not None:
                eachArea = eachArea + eachArea_list[i]
        #总面积大于6666.67平方米换算为单位亩
        if eachArea > 6666.67:
            eachArea = eachArea / 666.67
            return '%s亩' % str("{:.2f}".format(eachArea))
        return '%s平方米' % str("{:.2f}".format(eachArea))

    def print_projectname(self):
        file = self.e_inputLandFormFilePath.get()
        # 创建工作空间
        workbook = xlrd.open_workbook(file)
        # 选取第0个表
        sheet1 = workbook.sheet_by_index(0)
        porjectname = sheet1.cell(1, 8).value
        return porjectname

    # 地类分类表生成方法
    def landform_creat(self):
        file = self.e_inputLandFormFilePath.get()
        fileOutPath = self.e_outFilePath.get()
        projectName = self.print_projectname()
        # 存储数据表QSDWMC中去重后的名称
        list_q = []
        # 存储数据表Shape_Area数值
        sarea = []
        # 存储数据表DLMC名称
        dlmc = []
        # 存储数据表QSDWMC中未去重的名称
        qsdwmc = []

        land = ['水田', '水浇地', '旱地', '果园', '有林地', '沟渠',
                '道路', '坑塘水面', '设施农用地', '田坎', '仓储用地',
                '科教用地', '城市', '建制镇', '村庄', '公路', '水利设施用地',
                '墓葬地', '其他草地']
        # 需要输入数据的序列
        types = [2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 13, 14, 15, 16, 17, 18, 19, 20, 21]
        # 需要输入数据的序列
        types2 = ['B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T',
                  'U']

        # 读取数据表中的数据
        def readexcel():
            # 创建工作空间
            workbook = xlrd.open_workbook(file)
            # 选取第0个表
            sheet1 = workbook.sheet_by_index(0)
            # 获取第4列数据
            dlmclist = (sheet1.col_values(4))
            # 把第一行标题弹出去
            dlmclist.pop(0)
            # 获取第13列数据
            sarealist = (sheet1.col_values((13)))
            sarealist.pop(0)

            n1 = len(sarealist)
            # 数据添加到列表
            for a in range(0, n1):
                sarea.append(sarealist[a])

            n2 = len(dlmclist)
            for b in range(0, n2):
                dlmc.append(dlmclist[b])

            # 获取第5列数据
            iplist = (sheet1.col_values(5))
            # 把第一行标题弹出去
            iplist.pop(0)
            # 去掉重复的
            for i in iplist:
                if i not in list_q:
                    list_q.append(i)

            # 将有未去重的数据添加到列表
            for b in range(0, n2):
                qsdwmc.append(iplist[b])

        # 在模板中写入数据
        def writeexcel():
            file = '土地分类面积表2.xlsx'

            workbook = openpyxl.load_workbook(file)  # 创建工作空间
            sheet1 = workbook['Sheet1']  # 选取Sheet1表

            # 边框设置
            border_set = Border(left=Side(style='medium', color=colors.BLACK),
                                right=Side(style='medium', color=colors.BLACK),
                                top=Side(style='medium', color=colors.BLACK),
                                bottom=Side(style='medium', color=colors.BLACK))

            font1 = Font(name=u'Calibri', bold=False, italic=False, size=10.5)
            # 写入第一列
            n = len(list_q)
            for i in range(0, n):
                # 在第四行前插入一行，无格式
                workbook.active.insert_rows(4 + i)

                for j in range(1, 23):
                    # 设置插入行的单元格格式
                    sheet1.cell(row=4 + i, column=j).border = border_set
                    sheet1.cell(row=4 + i, column=j).font = font1
                    # 设置插入行的单元格数字格式为保留两位小数
                    sheet1.cell(row=4 + i, column=j).number_format = '0.00'
                    # 设置单元格自动换行，居中
                    sheet1.cell(row=4 + i, column=j).alignment = Alignment(horizontal='center', vertical='center',
                                                                           text_rotation=0, wrap_text=True,
                                                                           shrink_to_fit=False, indent=0)
                # 通过数据长度修改行高
                m = len(list_q[i])
                if m > 4:
                    sheet1.row_dimensions[4 + i].height = 30
                if m > 8:
                    sheet1.row_dimensions[4 + i].height = 45
                    # 写入单元格
                sheet1.cell(4 + i, 1).value = list_q[i]

            workbook.save('test1.xls')

            # 循环输入
            ldlen = len(land)
            llen = len(list_q)
            qlen = len(qsdwmc)
            qsdwmc.append('结束')
            # 确定行数
            for l1 in range(0, llen):
                # 循环，换行
                for ld1 in range(0, ldlen):
                    # 根据行标题写书行数据
                    if list_q[l1] == qsdwmc[0]:
                        # 遍历land
                        for ld1 in range(0, ldlen):
                            # 判断匹配
                            if land[ld1] == dlmc[0]:
                                # 用于判断单元格是否为空
                                data = sheet1.cell(4 + l1, types[ld1]).value
                                # 用于同名称数据相加
                                sum1 = sheet1.cell(4 + l1, 22).value
                                # 判断是否为空
                                if data is None:
                                    # 空值直接写入
                                    sheet1.cell(4 + l1, types[ld1]).value = sarea[0]
                                    workbook.save('test1.xlsx')
                                    # 写入后调整列宽
                                    # sheet1.column_dimensions[types2[ld1]].width = 8
                                    workbook.save('test1.xlsx')
                                else:
                                    # 非空值进行相加后写入
                                    sarea2 = data + sarea[0]
                                    sheet1.cell(4 + l1, types[ld1]).value = sarea2
                                    workbook.save('test1.xlsx')
                                # 判断合计列单元格是否为空
                                if sum1 is None:
                                    # 空值直接写入
                                    sheet1.cell(4 + l1, 22).value = sarea[0]
                                    workbook.save('test1.xlsx')
                                else:
                                    # 非空值进行相加后写入
                                    sum = sum1 + sarea[0]
                                    sheet1.cell(4 + l1, 22).value = sum
                                    workbook.save('test1.xlsx')
                            else:
                                pass
                        # 删除第0个，以保证每次循环取数正确
                        qsdwmc.pop(0)
                        dlmc.pop(0)
                        sarea.pop(0)
                    # 有效循环结束时，碰到空字符时结束全部循环
                    elif qsdwmc[0] == '':
                        return 0
                    # 判断不一样，跳过
                    else:

                        break
            # 保存
            workbook.save('test1.xlsx')

        # 计算合计部分
        def summation():
            file = 'test1.xlsx'
            # 创建工作空间
            workbook = xlrd.open_workbook(file)
            # 选取第0个表
            sheet1 = workbook.sheet_by_index(0)
            # 创建工作空间
            workbook2 = openpyxl.load_workbook(file)
            # 选取Sheet1表
            sheet2 = workbook2['Sheet1']
            # 在types中添加一列，以获取全部有效数值
            types.append(22)
            types2.append('V')
            # 设置列和动态变量
            cols_sum = locals()
            # 设置列值动态变量
            cols_list = locals()
            data_2 = []
            a = 0
            for i in types:
                # 创建动态变量列表
                cols_list['cols_list_' + str(i - 1)] = []
                # 获取i-1列的值
                cols_list['cols_list_' + str(i - 1)] = (sheet1.col_values(i - 1))
                # 删除第0个值，为了删除前三个标题
                cols_list['cols_list_' + str(i - 1)].pop(0)
                cols_list['cols_list_' + str(i - 1)].pop(0)
                cols_list['cols_list_' + str(i - 1)].pop(0)
                # 有效数值的行数
                lc = len(cols_list['cols_list_' + str(i - 1)])
                # 循环求和
                for j in cols_list['cols_list_' + str(i - 1)]:
                    # 用于检测单元格是否为空
                    data_1 = sheet2.cell(lc + 3, i).value
                    data_2 = re.findall(r'.+', str(data_1))

                    # 空值跳过
                    if j == '' or j ==' ':

                        pass
                    # 单元格为空值加入数值
                    elif data_1 is None or data_1 == ' ':

                        sheet2.cell(lc + 3, i).value = j

                    # 求和
                    else:

                        data = float(''.join(data_2))
                        j_1 = re.findall(r'.+', str(j))
                        j_2 = float(''.join(j_1))

                        sum2 = data + j_2


                        sheet2.cell(lc + 3, i).value = sum2
                        # float无法计算len，先变成str再取整数部分
                        sumlen = len(str(sum2).split('.')[0])
                        # 自适应列宽
                        sheet2.column_dimensions[types2[a]].width = sumlen + 3.5

                        workbook2.save('test2.xlsx')
                        workbook2.save(fileOutPath[:-5] + '土地分类表' + '.xlsx')

                a = a + 1

            workbook2.save('test2.xlsx')
            workbook2.save(fileOutPath[:-5] + '土地分类表' + '.xlsx')
            print(fileOutPath[:-5] + '土地分类表' + '.xlsx' + '完成')

        # 将.xlsx文件转换为.docx文件
        def xlsx2docx2():
            wb = load_workbook('test2.xlsx', data_only=False)
            document = Document()
            document.styles['Normal'].font.name = u'宋体'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            document.styles['Normal'].font.size = Pt(10)
            table_styles = ['Table Grid']
            p = document.add_paragraph()
            run = p.add_run('土地分类面积表')
            run.font.size = Pt(22)
            run.bold = True
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run2 = document.add_paragraph().add_run(
                '项目名称：%s                                                           平方米' % projectName)
            run2.font.size = Pt(12)

            section = document.sections[0]
            # 设置A4纸的宽度
            section.page_width = Cm(21)
            # 设置A4纸的高度
            section.page_height = Cm(29.7)
            # 设置页面横向
            section.orientation = WD_ORIENTATION.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            # 生成相应数量的空表
            ws = wb.worksheets[0]
            rows = list(ws.rows)
            table = document.add_table(rows=len(rows),
                                       cols=len(rows[0]),
                                       style=choice(table_styles))
            # 将表中的None值转换为0用于之后的选择需要
            for irow, row in enumerate(rows):
                for icol, col in enumerate(row):
                    table.cell(irow, icol).text = str(col.value)
                    # None全部转为0
                    if table.cell(irow, icol).text == 'None':
                        table.cell(irow, icol).text = '0'
            table.cell(1, 21).text = ' '
            table.cell(2, 21).text = ' '

            '''
            整数保留整数，小数保留两位小数
            删除会变成整数保留整数，小数保留三位小数
            也可改成全部保留两位小数
            '''
            # 将有用数值调整小数位数，其余清空
            for jrow in range(3, len(rows)):
                for jcol in range(1, len(rows[0])):
                    #table.cell(len(rows) - 2, 0).text = ' '
                    if table.cell(jrow, jcol).text != '0':
                        # 将保留3位小数的值调整为保留两位小数
                        if re.findall(r'(?=.{3})\.', (table.cell(jrow, jcol).text)) == ['.']:
                            table.cell(jrow, jcol).text = str("{:.2f}".format(float(table.cell(jrow, jcol).text)))
                    else:
                        table.cell(jrow, jcol).text = ' '
            # 合并单元格
            table.cell(0, 0).merge(table.cell(2, 0)).text = table.cell(0, 0).text
            table.cell(0, 1).merge(table.cell(0, 10)).text = table.cell(0, 1).text
            table.cell(0, 11).merge(table.cell(0, 19)).text = table.cell(0, 11).text
            table.cell(1, 1).merge(table.cell(1, 3)).text = table.cell(1, 1).text
            table.cell(1, 6).merge(table.cell(1, 10)).text = table.cell(1, 6).text
            table.cell(1, 11).merge(table.cell(1, 12)).text = table.cell(1, 11).text
            table.cell(1, 14).merge(table.cell(1, 16)).text = table.cell(1, 14).text
            # 调整表格总体宽度
            table.width = Cm(44)
            # 调整列宽
            for cell in table.columns[9].cells:
                cell.width = Cm(1.85)
            for cell in table.columns[13].cells:
                cell.width = Cm(1.85)
            for cell in table.columns[17].cells:
                cell.width = Cm(1.85)
            for cell in table.columns[18].cells:
                cell.width = Cm(1.85)
            for cell in table.columns[19].cells:
                cell.width = Cm(1.85)
            for cell in table.columns[20].cells:
                cell.width = Cm(1.85)

            for table_row in range(0, len(rows)):
                for table_col in range(0, 22):
                    # 垂直居中
                    table.cell(table_row, table_col).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.cell(table_row, table_col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for row in table.rows:
                # 设置行高
                row.height = Cm(0.65)

            document.save('report2.docx')
            document.save(fileOutPath[:-5] + '土地分类表' + '.docx')
            print(fileOutPath[:-5] + '土地分类表' + '.docx')

        if __name__ == '__main__':
            readexcel()
            writeexcel()
            summation()
            xlsx2docx2()

    # 界址点成果表
    def creat_JZDForm(self):
        #创建基础值
        face = []
        jlist = []
        pointx = []
        pointy = []
        rowlist = []
        mj = []
        mjm = []
        file = self.e_inputJZDFormFilePath.get()

        # 读取导入的界址点信息表
        def readJZD():
            # 创建工作空间
            workbook = xlrd.open_workbook(file)
            # 选取第0个表
            sheet1 = workbook.sheet_by_index(0)

            # 获取值
            mj1 = (sheet1.col_values(1))
            mj1.pop(0)

            mjm1 = (sheet1.col_values(2))
            mjm1.pop(0)

            face1 = (sheet1.col_values((3)))
            # 把第一行标题弹出去
            face1.pop(0)

            jlist1 = (sheet1.col_values(4))
            jlist1.pop(0)

            pointx1 = (sheet1.col_values((5)))
            pointx1.pop(0)

            pointy1 = (sheet1.col_values((6)))
            pointy1.pop(0)

            # 添加数据
            n1 = len(face1)
            for a1 in range(0, n1):
                face.append(int(face1[a1]))

            n2 = len(jlist1)
            for a2 in range(0, n2):
                jlist.append(jlist1[a2])
                rowlist.append(int(jlist1[a2][1:]))

            n3 = len(pointx1)
            for a3 in range(0, n3):
                pointx.append(pointx1[a3])

            n4 = len(pointy1)
            for a4 in range(0, n4):
                pointy.append(pointy1[a4])

            n5 = len(mj1)
            for a5 in range(0, n5):
                mj.append(mj1[a5])

            n6 = len(mjm1)
            for a6 in range(0, n6):
                mjm.append(mjm1[a6])

        # 生成界址点成果表.xlsx
        def writeJZD():
            wb = Workbook()

            # 边框设置
            border_set = Border(left=Side(style='medium', color=colors.BLACK),
                                right=Side(style='medium', color=colors.BLACK),
                                top=Side(style='medium', color=colors.BLACK),
                                bottom=Side(style='medium', color=colors.BLACK))
            # 字体设置
            font1 = Font(name=u'Calibri', bold=False, italic=False, size=10.5)

            n = 0
            ws = locals()
            for i in jlist:

                # 通过J1判断，生成多个地区的表
                if i == 'J1':
                    ws['ws' + str(face[n])] = wb.create_sheet(index=face[n], title='Sheet' + str(face[n]))
                    # 写入表头和样式
                    ws['ws' + str(face[n])].cell(row=1, column=1).value = '点号'
                    ws['ws' + str(face[n])].cell(row=1, column=2).value = 'X坐标'
                    ws['ws' + str(face[n])].cell(row=1, column=3).value = 'Y坐标'
                    ws['ws' + str(face[n])].cell(row=1, column=4).value = '面积'
                    ws['ws' + str(face[n])].cell(row=1, column=5).value = '面积亩'

                    ws['ws' + str(face[n])].column_dimensions['B'].width = 20
                    ws['ws' + str(face[n])].column_dimensions['C'].width = 20
                    ws['ws' + str(face[n])].column_dimensions['D'].width = 20
                    ws['ws' + str(face[n])].column_dimensions['E'].width = 20
                    # 设置格式
                    for i1 in range(1, 6):
                        ws['ws' + str(face[n])].cell(row=1, column=i1).border = border_set
                        ws['ws' + str(face[n])].cell(row=1, column=i1).font = font1
                        ws['ws' + str(face[n])].cell(row=1, column=i1).alignment = Alignment(horizontal='center',
                                                                                                  vertical='center',
                                                                                                  text_rotation=0,
                                                                                                  wrap_text=True,
                                                                                                  shrink_to_fit=False,
                                                                                                  indent=0)
                n = n + 1

            # 写入数据和样式
            n1 = 0
            for j in face:

                # 确定表，并在每个表写入数据
                if j:
                    sheet1 = wb['Sheet{0}'.format(j)]
                    sheet1.cell(row=int(rowlist[n1]) + 1, column=1).value = jlist[n1]
                    sheet1.cell(row=int(rowlist[n1]) + 1, column=2).value = pointy[n1]
                    sheet1.cell(row=int(rowlist[n1]) + 1, column=3).value = pointx[n1]
                    sheet1.cell(row=int(rowlist[n1]) + 1, column=4).value = mj[n1]
                    sheet1.cell(row=int(rowlist[n1]) + 1, column=5).value = mjm[n1]

                    wb.save('test.xlsx')

                    for j1 in range(1, 6):
                        # 设置插入行的单元格格式
                        sheet1.cell(row=int(rowlist[n1]) + 1, column=j1).border = border_set
                        # 设置单元格字体
                        sheet1.cell(row=int(rowlist[n1]) + 1, column=j1).font = font1
                        # 设置插入行的单元格数字格式为保留三位小数
                        sheet1.cell(row=int(rowlist[n1]) + 1, column=j1).number_format = '0.000'
                        # 单元格居中
                        sheet1.cell(row=int(rowlist[n1]) + 1, column=j1).alignment = Alignment(horizontal='center',
                                                                                                    vertical='center',
                                                                                                    text_rotation=0,
                                                                                                    wrap_text=True,
                                                                                                    shrink_to_fit=False,
                                                                                                    indent=0)

                n1 = n1 + 1
            # 遍历所有表，修改每个表最后改成J1
            sheets = wb.sheetnames
            for j1 in range(1, len(sheets)):
                sheet = wb[sheets[j1]]
                sheet.cell(row=sheet.max_row, column=1).value = 'J1'

            # 删除多余表
            ws = wb['Sheet']
            wb.remove(ws)
            wb.save('test.xlsx')

        if __name__ == '__main__':
            readJZD()
            writeJZD()



def main():
    Aotu = Aotumation()
    Aotu.gui_arrang()
    mainloop()
    pass


if __name__ == "__main__":
    main()
